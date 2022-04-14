# -*- coding: utf-8 -*-
import sqlite3
import sys

import docx
import numpy
import pandas
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from pathlib import Path


# factory for returning list of dicts against list of tuples from sql result (smth like dict cursor)
def dict_factory(cursor, row):
    dict_obj = {}
    for idx, col in enumerate(cursor.description):
        dict_obj[col[0]] = row[idx]
    return dict_obj


# computing cagr for given dataframe and factor, ignoring nan
def compute_cagr(df, factor):
    year_cols = df.columns.levels[1]
    year1 = None
    year2 = None
    cagr = None
    try:
        year1 = next(year for year in year_cols if not numpy.isnan(df[6, year].values[0]))
        year2 = next(year for year in reversed(year_cols) if not numpy.isnan(df[6, year].values[0]))
    except StopIteration:
        pass

    if None not in (year1, year2):
        cagr = (df[factor, year2].values[0] / df[factor, year1].values[0]) ** (1.0 / (year2 - year1)) - 1
    return cagr, year1, year2


# init data
db_path = Path(sys.argv[0]).parent.joinpath('test_task_python', 'test.db')

# connecting to db
with sqlite3.connect(str(db_path)) as conn:
    conn.row_factory = dict_factory
    curs = conn.cursor()
    query = """
        SELECT factor, year, country, res FROM `testidprod`
        WHERE partner IS NULL
            AND state IS NULL
            AND bs = 0
            AND factor IN (1, 2)
    """
    curs.execute(query)

sql_result = curs.fetchall()

# There is a dict building process with dict and tuple comprehensions
df_dict = {
    (factor, year):
        sum(item['res'] if item['res'] is not None else 0 for item in sql_result
            if item['year'] == year and item['factor'] == factor)
        if year in set(item['year'] for item in sql_result if item['factor'] == factor) else numpy.NaN
    for year in range(2006, 2021)
    for factor in set(item['factor'] for item in sql_result)
}

# There is a dict building process with traditional "for" loops
# Result is similar with dict comprehension approach
# df_dict = {}
# factors = set(item['factor'] for item in sql_result)
# for factor in factors:
#     years = set(item['year'] for item in sql_result if item['factor'] == factor)
#     for year in range(2006, 2021):
#         if year in years:
#             all_res = (
#                 item['res'] if item['res'] is not None else 0 for item in sql_result
#                 if item['year'] == year and item['factor'] == factor
#             )
#             df_dict[(factor, year)] = sum(all_res)
#         else:
#             df_dict[(factor, year)] = numpy.NaN

# creating dataframe
df = pandas.DataFrame(df_dict, index=('world',)).rename_axis(('Factor', 'Year'), axis=1)

# adding new column with factor 6
df = df.join(df[[2]].div(df[1], level=1).rename(columns={2: 6}))

# Exporting dataframe to Excel document
df.to_excel('report.xlsx')

# exporting transposed dataframe to Word document
df_t = df[[6]].transpose().reset_index().rename(columns={'world': 'World Value'})
doc = docx.Document()

# table adding
table = doc.add_table(df_t.shape[0] + 1, df_t.shape[1])

# merging first column cells
table.columns[0].cells[1].merge(table.columns[0].cells[-1])

for i in range(df_t.shape[-1]):
    table.cell(0, i).text = str(df_t.columns[i])
for i in range(df_t.shape[0]):
    for j in range(1, df_t.shape[-1]):
        value = df_t.values[i, j]
        if j == 1:
            value = int(value)
        else:
            value = round(float(value), 2)
        table.cell(i + 1, j).text = str(value)

# applying alignment and adding factor
table.columns[0].cells[1].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
p = table.columns[0].cells[1].add_paragraph(str(int(df_t.values[0, 0])))
p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

# Adding cagr information
cagr, year1, year2 = compute_cagr(df, 6)
if None not in (cagr, year1, year2):
    doc.add_paragraph('Factor 6 {} by avg {}% every year from {} to {}.'.format(
        'grew' if cagr > 0 else 'decreased', round(cagr * 100, 2), year1, year2
    ))

doc.save('report.docx')
